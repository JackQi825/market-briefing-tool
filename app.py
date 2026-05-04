import re
import os
import zipfile
import hashlib
import json
from io import BytesIO
from html import escape
from pathlib import Path
from urllib.parse import urlparse

import fitz
import requests
import streamlit as st
import streamlit.components.v1 as components
from bs4 import BeautifulSoup
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from PIL import Image, UnidentifiedImageError
from pypdf import PdfReader


APP_DIR = Path(__file__).resolve().parent
load_dotenv(APP_DIR / ".env")

APP_DISPLAY_NAME = "Jack 市场沟通助手"

st.set_page_config(page_title=APP_DISPLAY_NAME, page_icon="📊", layout="wide")

DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DEFAULT_DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-v4-flash")
APP_PASSWORD = os.getenv("APP_PASSWORD", "")
MAX_INPUT_CHARS = 50000
MAX_COMPARE_DOCUMENTS = 5
MAX_EXTRACTED_IMAGES = 8
MIN_CHART_WIDTH = 420
MIN_CHART_HEIGHT = 240

OUTPUT_STYLES = {
    "略懂一二客户版": """
受众画像：客户平时会看财经新闻，知道美联储、降息、美元、港股、美股、黄金等基本概念，但不一定能系统理解资产配置。
写作目标：让客户觉得“这个观点有判断、有启发”，不是普通市场新闻复述。
语言要求：专业但不端着，可以使用少量金融词汇，但每个关键判断都要翻译成客户能理解的话。
重点输出：
- 先讲清楚市场主线，再讲对客户组合意味着什么。
- 资产观点要有明确方向，例如偏积极、偏中性、需要控制仓位，而不是含糊其辞。
- 可以适度讲短期机会，但必须说明适合什么客户、风险边界在哪里。
禁止：
- 不要写成研报摘要。
- 不要解释过于基础的概念，避免让客户觉得被“科普”。
- 不要堆砌术语。
""",
    "完全小白解释版": """
受众画像：客户几乎不看市场，对股票、债券、汇率、利率、黄金之间的关系不熟悉。
写作目标：把复杂市场翻译成“听得懂、知道和自己有什么关系、知道下一步可以关注什么”。
语言要求：尽量少用术语；必须使用类比、生活化表达或一句话解释专业词。
重点输出：
- 每出现一个重要金融概念，都要用括号或短句解释，例如“降息，就是资金价格变便宜”。
- 先回答客户最关心的问题：市场为什么动、我的钱会受什么影响、现在该做什么准备。
- 各资产观点要多讲“为什么”，少讲复杂指标。
- 建议动作要具体，例如检查现金比例、确认持有期限、分清核心配置和机会型配置。
禁止：
- 不要假设客户懂宏观逻辑。
- 不要使用“久期、风险偏好、收益率曲线、估值修复”等词后不解释。
- 不要给客户制造焦虑。
""",
    "客户经理自用版": """
受众画像：财富管理客户经理本人，用于快速理解材料、准备客户沟通、筛选可跟进客户。
写作目标：帮助客户经理快速抓住主线、机会、风险和可转化成话术的卖点。
语言要求：结构清晰、判断直接、信息密度高，可以保留必要专业表达，但每个观点要能转成客户语言。
重点输出：
- 明确哪些是主线观点，哪些是短期机会，哪些只是观察信号。
- 标注适合跟进的客户类型，例如现金仓位高、美元资产多、港股关注度高、债券配置不足、风险偏好较高等。
- 给出客户沟通角度：怎么开口、怎么解释、怎么提醒风险。
- 对材料里的图表、数据、分歧点、隐含机会要敏感。
禁止：
- 不要只写宏观大段分析。
- 不要只给结论不给客户跟进动作。
- 不要把机会型观点写成确定性推荐。
""",
    "内部群同事分享版": """
受众画像：同事、团队内部群、晨会快速分享，大家懂基本金融语言，但时间有限。
写作目标：让同事在1-2分钟内看懂“这份材料最值得转发/转述的观点是什么”。
语言要求：短、硬、清楚，适合微信群或内部群直接发；可以使用项目符号和加粗式表达，但不要冗长。
重点输出：
- 开头先给一句最核心结论。
- 后面用3-5条列出关键判断，每条都带“为什么”。
- 单独列“可对客户讲的话”和“需要谨慎表达的点”。
- 对短期机会要提炼成同事能直接理解的跟进方向。
禁止：
- 不要写成长篇报告。
- 不要写客套话。
- 不要重复材料原文的大段表述。
""",
}

TOUCH_COPY_SYSTEM_PROMPT = """
你是一位资深的财富管理客户经理触达文案专家。

你的任务:用户会提供市场观点/研报/新闻,你需要将其转化为 3种不同风格的客户触达文案。

【核心原则 — 3S 法则】
1. Selective(精准):内容要有针对性,不要平庸的"市场总结"
2. Sharp(尖锐):必须给出明确判断,不要模糊的"投资者保持谨慎"
3. Single(单一):一条文案只讲 1 个核心观点,不要堆积信息

【写作公式】
每条文案必须包含三要素:
- 【市场事件】发生了什么(简洁)
- 【对客户的影响】这意味着什么(专业判断)
- 【建议动作】客户可以做什么(具体可执行)

【绝对禁忌】
- ❌ 不要出现"投资有风险,入市需谨慎"等套话
- ❌ 不要照抄原文,要重新组织语言
- ❌ 不要用过多金融术语,用客户能懂的话
- ❌ 不要超过指定字数

【输出格式】
严格按以下 JSON 格式输出,不要输出 Markdown,不要输出代码块:
{
  "one_liner": "一句话市场观点的内容",
  "deep_push": "针对性深度推送的内容",
  "care_touch": "钩子型触达的内容"
}

【字数要求】
- one_liner: 30-80 字
- deep_push: 150-250 字
- care_touch: 50-100 字

补充要求:
- one_liner: 一段话,带 emoji 装饰,必须包含发生了什么 + 对客户意味着什么 + 1 个具体建议,适合周一群发。
- deep_push: 称呼 + 3 段落,必须包含事实 + 判断 + 建议 + 行动召唤,适合配置缺口客户、不同地域市场或板块关注客户。
- care_touch: 这是“钩子型触达”,目标是让不关心市场的客户也想回复你。不要写“保重身体”“调整心态”“市场波动较大请保持平常心”这类温暖空话。
- care_touch 必须从以下 4 种模式里选择一种来写,但不要在文案里标注模式名称:
  1. 反常识: 用一个和客户直觉相反的市场观察开头,例如“这轮上涨最值得看的可能不是涨得最多的资产”。
  2. 场景代入: 把市场变化放进客户真实生活或资产配置场景里,例如“如果这笔钱 3 个月内不用,现在最该看的不是收益率高低,而是流动性和波动”。
  3. 冷知识: 提供一个有信息密度的小观察,例如利率、美元、黄金、港股、债券之间的联动,但不要编造具体数据。
  4. 轻幽默: 语气轻一点,但保持专业,可以有一点反差感,不能油腻、不能像段子。
- care_touch 必须包含一个“可回复的钩子”,例如“您最近更关心稳一点还是机会多一点?”、“我可以帮您看一眼组合里有没有受影响的部分。”
- care_touch 不能有任何销售话术,不能直接推荐产品,不能催促购买,不能承诺收益。
- 可以基于用户资料和你自身通用金融市场知识共同分析和补充,但不要编造具体数据。
"""

CHART_RELEVANCE_KEYWORDS = [
    "图",
    "图表",
    "表",
    "数据",
    "走势",
    "趋势",
    "比较",
    "对比",
    "历史",
    "区间",
    "分位",
    "估值",
    "市盈率",
    "市净率",
    "收益率",
    "曲线",
    "利差",
    "通胀",
    "CPI",
    "PCE",
    "PMI",
    "就业",
    "盈利",
    "EPS",
    "增长",
    "美元",
    "汇率",
    "黄金",
    "原油",
    "铜",
    "资产表现",
    "回报",
    "波动率",
    "资金流",
    "配置",
    "超配",
    "低配",
    "机会",
    "风险",
]

REQUEST_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0 Safari/537.36"
    )
}


def inject_custom_css():
    st.markdown(
        """
        <style>
        :root {
            --app-bg: #ffffff;
            --panel: #ffffff;
            --panel-soft: #fff7f7;
            --ink: #1f1f1f;
            --muted: #666666;
            --line: #e7e7e7;
            --accent: #db0011;
            --accent-strong: #a8000d;
            --accent-soft: #fff1f2;
            --dark: #20242a;
        }

        .stApp {
            background:
                linear-gradient(180deg, #ffffff 0%, #ffffff 54%, #fafafa 100%);
            color: var(--ink);
        }

        .block-container {
            max-width: 1180px;
            padding-top: 28px;
            padding-bottom: 48px;
        }

        section[data-testid="stSidebar"] {
            background: #fafafa;
            border-right: 1px solid var(--line);
        }

        .app-hero {
            background:
                linear-gradient(90deg, rgba(219, 0, 17, 0.08), rgba(255, 255, 255, 0) 38%),
                var(--panel);
            border: 1px solid var(--line);
            border-top: 5px solid var(--accent);
            border-radius: 8px;
            padding: 24px 26px;
            margin-bottom: 20px;
            box-shadow: 0 12px 28px rgba(0, 0, 0, 0.06);
        }

        .app-kicker {
            color: var(--accent);
            font-size: 13px;
            font-weight: 700;
            letter-spacing: 0;
            margin-bottom: 8px;
        }

        .app-title {
            color: var(--ink);
            font-size: 36px;
            font-weight: 760;
            line-height: 1.16;
            margin: 0 0 8px;
        }

        .app-subtitle {
            color: var(--muted);
            font-size: 15px;
            line-height: 1.7;
            margin: 0;
        }

        .status-row {
            display: flex;
            gap: 10px;
            flex-wrap: wrap;
            margin-top: 16px;
        }

        .status-pill {
            display: inline-flex;
            align-items: center;
            border: 1px solid var(--line);
            background: #ffffff;
            border-radius: 999px;
            padding: 7px 11px;
            color: var(--muted);
            font-size: 13px;
            font-weight: 600;
        }

        .status-pill.ready {
            border-color: rgba(219, 0, 17, 0.24);
            color: var(--accent-strong);
            background: var(--accent-soft);
        }

        .status-pill.warn {
            border-color: rgba(122, 122, 122, 0.32);
            color: #555555;
            background: #f7f7f7;
        }

        div[data-testid="stFileUploader"],
        div[data-testid="stTextArea"],
        div[data-testid="stTextInput"] {
            background: var(--panel);
            border: 1px solid var(--line);
            border-radius: 8px;
            padding: 14px 14px 8px;
            box-shadow: 0 6px 18px rgba(0, 0, 0, 0.035);
        }

        div[data-testid="stTextInput"] input,
        div[data-testid="stTextArea"] textarea {
            border-radius: 6px;
        }

        .stButton > button {
            background: var(--accent);
            color: #ffffff;
            border: 1px solid var(--accent);
            border-radius: 7px;
            min-height: 46px;
            font-weight: 700;
            box-shadow: 0 8px 18px rgba(219, 0, 17, 0.16);
        }

        .stButton > button:hover {
            background: var(--accent-strong);
            border-color: var(--accent-strong);
            color: #ffffff;
        }

        h1, h2, h3 {
            letter-spacing: 0;
        }

        h2 {
            padding-top: 8px;
            color: var(--ink);
            border-left: 4px solid var(--accent);
            padding-left: 10px;
        }

        .stTabs [data-baseweb="tab-list"] {
            gap: 8px;
            border-bottom: 1px solid var(--line);
        }

        .stTabs [data-baseweb="tab"] {
            border-radius: 6px 6px 0 0;
            padding: 10px 14px;
            font-weight: 700;
        }

        div[data-testid="stExpander"] {
            background: var(--panel);
            border: 1px solid var(--line);
            border-radius: 8px;
        }

        hr {
            border-color: var(--line);
        }

        div[data-testid="stAlert"] {
            border-radius: 8px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


inject_custom_css()


def get_deepseek_client():
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        return None
    return OpenAI(api_key=api_key, base_url=DEEPSEEK_BASE_URL)


def require_password():
    if not APP_PASSWORD:
        return True

    if st.session_state.get("authenticated"):
        return True

    st.markdown(
        """
        <div class="app-hero">
            <div class="app-kicker">财富管理客户沟通工作台</div>
            <h1 class="app-title">Jack 市场沟通助手</h1>
            <p class="app-subtitle">请输入访问密码，进入市场观点分析工具。</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    password = st.text_input("访问密码", type="password")
    if st.button("进入工具", type="primary"):
        if password == APP_PASSWORD:
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("密码不正确，请重试。")
    return False


if not require_password():
    st.stop()


def format_documents_for_prompt(documents):
    if not documents:
        return ""

    per_document_limit = max(4000, MAX_INPUT_CHARS // max(len(documents), 1))
    blocks = []
    for index, document in enumerate(documents, start=1):
        content = clean_text(document["text"])[:per_document_limit]
        blocks.append(
            f"【材料{index}：{document['name']}】\n{content}"
        )
    return "\n\n---\n\n".join(blocks)[:MAX_INPUT_CHARS]


def build_deepseek_prompt(text, output_style, documents=None):
    documents = documents or []
    trimmed_text = format_documents_for_prompt(documents) if documents else text[:MAX_INPUT_CHARS]
    document_count = len(documents) if documents else 1
    compare_instruction = ""
    if document_count >= 2:
        compare_instruction = f"""

特别说明：本次输入包含 {document_count} 份材料，请先分别理解每份材料，再做交叉分析。不要把多份材料简单拼接成一份来看。

九、多文档对比分析
请输出以下内容：
- 共识观点：列出3-6条多份材料都支持或方向一致的观点。每条说明哪些材料支持、为什么重要。
- 分歧点：列出2-5条不同材料之间的判断差异。每条说明分歧来自哪里、可能原因是什么、客户经理应该如何解释。
- 强弱信号排序：把最确定、最有一致性的市场信号排在前面，把证据不足或分歧较大的观点排在后面。
- 对客户沟通的取舍建议：哪些观点可以作为主线讲，哪些只能作为补充观察，哪些需要谨慎表达。
- 如果某份材料明显更偏短期交易、某份更偏中长期配置，请明确区分。
"""

    style_instruction = OUTPUT_STYLES.get(output_style, OUTPUT_STYLES["略懂一二客户版"])
    return f"""
你是一名财富管理客户经理的市场研究助理。请基于以下市场展望文件内容，生成一份“客户经理可用版市场观点分析”。

本次输出风格：{output_style}
风格要求：{style_instruction}

请严格按照以下结构输出，不要遗漏任何部分：

一、文档核心观点
- 用一句话总结文档主线。
- 提炼3-5个最重要观点。
- 每个观点必须说明“为什么”。

二、当前市场主线
- 说明当前市场处在什么环境，例如风险偏好、利率周期、美元趋势、通胀、企业盈利、政策预期。
- 不要只复述原文，要翻译成客户能听懂的话。

三、各资产大类观点及支撑逻辑
请分别输出股票、债券、外汇、商品、现金/货币基金：
- 未来观点：看多 / 中性 / 谨慎 / 看空。
- 重点关注方向或区域。
- 支撑逻辑：至少3条，每条都解释为什么。
- 主要风险：至少2条，每条都解释为什么。
- 客户可理解的解释：用简单语言说清楚。

四、短期机会型投资观点
请主动识别原文中是否提到未来1-3个月可以关注的机会型、战术型、阶段性配置观点。
可以包括但不限于：
- 某个区域市场机会，例如港股、日本股、美股科技、亚洲资产、中国资产等。
- 某个行业或主题机会，例如AI、红利、高股息、科技、消费、资源品、制造业、金融等。
- 某类债券机会，例如长久期、短久期、投资级债、高收益债、利率债、信用债等。
- 某类商品或外汇机会，例如黄金、铜、原油、美元、日元、人民币等。
- 事件驱动机会，例如降息交易、政策催化、财报季、地缘风险、再通胀交易等。

输出要求：
- 如果原文有明确机会，请列出3-6条。
- 每条都按“机会是什么 / 为什么现在值得关注 / 适合什么客户 / 需要注意什么风险”的格式写。
- 必须区分“核心配置”和“机会型配置”，不要把短期机会写成确定性推荐。
- 如果原文没有明确提到短期机会，不要硬编，写“原文没有给出明确短期机会型观点”，再说明可以继续观察哪些信号。

五、文档图表关键信息
请特别关注原文里的图表标题、图表注释、表格、数据对比、走势图描述、分位数、历史区间、环比/同比变化、资产表现排序等信息。
输出要求：
- 如果原文文字里能读到图表或表格信息，请列出2-5条“图表传递的关键信号”。
- 每条都按“图表/数据说明了什么 / 对市场判断有什么帮助 / 客户怎么理解”的格式写。
- 如果文档包含图片图表但正文没有足够图注或数据，请写“图表图片需要人工核对”，并尽量提示客户经理应该回到原文查看哪类图表、哪一段附近的图表或哪组数据。
- 不要编造图表中没有出现的具体数字。

六、主要风险
- 列出3-5条。
- 每条说明为什么会影响市场或客户组合。

七、五月市场展望微信版
要求：
- 100-200字。
- 像客户经理发给客户的自然市场回顾，不要像研报。
- 包括近期发生了什么、为什么会这样、五月怎么看、对配置有什么启发。
- 如果原文有短期机会型观点，可以自然提一句“阶段性可以关注什么”，但不要写成买卖建议。
- 不承诺收益，不使用“必涨”“一定”“稳赚”等表达。

八、1分钟电话汇报版
要求：
- 开头自然。
- 逻辑清楚。
- 能解释市场主线。
- 能自然引导客户关注资产配置。
- 如果原文有机会型观点，要自然说明“哪些可以作为机会型观察，不适合当成重仓押注”。
- 不要太销售。

写作要求：
1. 不要空泛总结。
2. 每一个核心观点后面都必须说明“为什么”。
3. 不允许只复述原文，要转化成客户能听懂的话。
4. 对专业词汇要自动翻译成通俗解释。
5. 微信版和电话版必须像真实客户经理会说的话。
6. 如果原文没有提到某类资产，请明确写“原文信息有限”，并基于宏观线索谨慎推断，不要编造具体数据。
7. 对机会型观点要特别敏感，看到“短期、战术、阶段性、关注、受益、催化、交易、主题、反弹、修复、窗口期、配置价值”等表达时，要单独提炼到第四部分。
8. 任何机会型观点都必须配风险提示，不能表达成确定收益或直接买卖建议。
9. 如果输入包含多份材料，请明确标注观点来自哪份材料，并区分“多份材料共识”和“单一材料观点”。
{compare_instruction}

以下是市场展望文件内容：
---
{trimmed_text}
---
"""


def call_deepseek_market_analysis(text, model, output_style, documents=None):
    client = get_deepseek_client()
    if client is None:
        raise ValueError("没有读取到 DEEPSEEK_API_KEY。请先在 .env 文件里填写你的 DeepSeek API Key。")

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "你是一名财富管理市场研究助理。"},
            {"role": "user", "content": build_deepseek_prompt(text, output_style, documents)},
        ],
        temperature=0.3,
    )
    return response.choices[0].message.content


def parse_touch_copy_response(raw_text):
    cleaned = raw_text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError(f"DeepSeek 没有返回可识别的 JSON：{raw_text[:200]}")

    json_text = re.sub(r",\s*}", "}", cleaned[start:end + 1])
    try:
        data = json.loads(json_text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"DeepSeek 返回的 JSON 格式不完整：{exc}") from exc

    required_keys = ("one_liner", "deep_push", "care_touch")
    missing_keys = [key for key in required_keys if not data.get(key)]
    if missing_keys:
        raise ValueError(f"DeepSeek 返回内容缺少字段：{', '.join(missing_keys)}")

    return {key: str(data[key]).strip() for key in required_keys}


def call_deepseek_touch_copy(text, model):
    client = get_deepseek_client()
    if client is None:
        raise ValueError("没有读取到 DEEPSEEK_API_KEY。请先在 .env 文件里填写你的 DeepSeek API Key。")

    trimmed_text = text[:MAX_INPUT_CHARS]
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": TOUCH_COPY_SYSTEM_PROMPT},
            {"role": "user", "content": trimmed_text},
        ],
        temperature=0.4,
    )
    return parse_touch_copy_response(response.choices[0].message.content)


def call_deepseek_followup(question, model):
    client = get_deepseek_client()
    if client is None:
        raise ValueError("没有读取到 DEEPSEEK_API_KEY。请先在 .env 文件里填写你的 DeepSeek API Key。")

    source_text = st.session_state.get("analysis_source_text", "")
    report = st.session_state.get("analysis_result", "")
    source_text = source_text[:MAX_INPUT_CHARS]

    messages = [
        {
            "role": "system",
            "content": (
                "你是一名财富管理市场研究助理。请结合用户上传的市场材料、已经生成的市场观点分析，"
                "以及你自身的通用金融市场知识和推理能力，回答用户追问。上传材料是优先依据，"
                "但不要被材料完全限制；当材料没有覆盖问题时，可以给出模型补充判断。"
                "回答时要清楚区分：1）原文依据；2）模型补充判断；3）需要进一步核实的信息。"
                "不要编造原文没有的具体数据或假装原文提过；涉及投资观点时必须说明依据和风险边界；"
                "不承诺收益，不给确定性买卖建议。"
            ),
        },
        {
            "role": "user",
            "content": (
                "以下是本次市场材料原文摘要/正文：\n"
                f"{source_text}\n\n"
                "以下是已经生成的市场观点分析：\n"
                f"{report}\n\n"
                "请记住以上上下文。后续回答追问时，优先参考本次材料；如需超出材料范围，"
                "可以使用你的通用市场知识进行补充，但必须明确写出“模型补充判断”，并提示不等同于原文结论。"
            ),
        },
    ]

    for message in st.session_state.get("followup_messages", [])[-8:]:
        messages.append({"role": message["role"], "content": message["content"]})
    messages.append({"role": "user", "content": question})

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.3,
    )
    return response.choices[0].message.content


SECTION_KEYWORDS = {
    "股票": [
        "股票",
        "股市",
        "权益",
        "A股",
        "港股",
        "美股",
        "纳指",
        "标普",
        "道指",
        "沪指",
        "创业板",
        "科技股",
        "估值",
        "盈利",
    ],
    "债券": [
        "债券",
        "债市",
        "国债",
        "美债",
        "收益率",
        "利率",
        "久期",
        "信用债",
        "利差",
        "降息",
        "加息",
        "央行",
    ],
    "外汇": [
        "外汇",
        "汇率",
        "美元",
        "人民币",
        "欧元",
        "日元",
        "英镑",
        "美指",
        "美元指数",
        "贬值",
        "升值",
    ],
    "商品": [
        "商品",
        "原油",
        "黄金",
        "白银",
        "铜",
        "铁矿",
        "煤炭",
        "农产品",
        "大宗",
        "能源",
        "贵金属",
    ],
    "现金/货币基金": [
        "现金",
        "货币基金",
        "货基",
        "存款",
        "短端",
        "流动性",
        "闲钱",
        "短期资金",
    ],
}

RISK_KEYWORDS = [
    "风险",
    "不确定",
    "波动",
    "下行",
    "压力",
    "衰退",
    "通胀",
    "地缘",
    "政策",
    "流动性",
    "违约",
    "回撤",
]

CORE_KEYWORDS = [
    "市场",
    "经济",
    "宏观",
    "政策",
    "通胀",
    "就业",
    "利率",
    "降息",
    "加息",
    "增长",
    "风险偏好",
    "流动性",
]

ASSET_CONFIG = {
    "股票": {
        "keywords": SECTION_KEYWORDS["股票"],
        "sub_assets": ["美国", "欧洲", "日本", "亚洲", "中国", "港股", "A股", "美股"],
        "default_focus": "重点关注报告中提到的盈利改善、估值位置和政策支持区域。",
    },
    "债券": {
        "keywords": SECTION_KEYWORDS["债券"],
        "sub_assets": ["国债", "美债", "信用债", "投资级债", "高收益债", "短久期", "长久期"],
        "default_focus": "重点关注利率方向、降息节奏和收益率曲线变化。",
    },
    "外汇": {
        "keywords": SECTION_KEYWORDS["外汇"],
        "sub_assets": ["美元", "人民币", "日元", "欧元", "英镑"],
        "default_focus": "重点关注美元方向、主要央行政策差异和人民币汇率压力。",
    },
    "商品": {
        "keywords": SECTION_KEYWORDS["商品"],
        "sub_assets": ["黄金", "原油", "铜", "白银", "能源", "贵金属"],
        "default_focus": "重点关注避险需求、供需变化和全球增长预期。",
    },
    "现金/货币基金": {
        "keywords": SECTION_KEYWORDS["现金/货币基金"],
        "sub_assets": ["现金", "货币基金", "短期存款", "短端理财"],
        "default_focus": "适合作为流动性和等待机会的短期资金安排。",
    },
}

POSITIVE_WORDS = [
    "看好",
    "上行",
    "上涨",
    "改善",
    "修复",
    "回升",
    "机会",
    "支撑",
    "受益",
    "超配",
    "增配",
    "流入",
    "宽松",
    "降息",
    "盈利增长",
]

NEGATIVE_WORDS = [
    "看空",
    "下行",
    "下跌",
    "回落",
    "压力",
    "承压",
    "风险",
    "波动",
    "衰退",
    "低配",
    "减配",
    "收紧",
    "高估",
    "不确定",
]

ENVIRONMENT_LABELS = {
    "风险偏好": ["风险偏好", "避险", "情绪", "波动"],
    "利率周期": ["利率", "降息", "加息", "收益率", "央行", "宽松"],
    "美元趋势": ["美元", "美元指数", "美指", "汇率"],
    "通胀": ["通胀", "CPI", "PCE", "物价"],
    "企业盈利": ["盈利", "利润", "业绩", "EPS"],
    "政策预期": ["政策", "财政", "货币政策", "监管", "刺激"],
}


def clean_text(text):
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_sentences(text):
    text = clean_text(text)
    parts = re.split(r"(?<=[。！？!?；;])\s*|\n+", text)
    sentences = [part.strip() for part in parts if len(part.strip()) >= 8]
    return sentences


def score_sentence(sentence, keywords):
    score = 0
    for keyword in keywords:
        if keyword.lower() in sentence.lower():
            score += 2
    if re.search(r"\d+(\.\d+)?%?", sentence):
        score += 1
    if 20 <= len(sentence) <= 120:
        score += 1
    return score


def pick_sentences(sentences, keywords, limit=3):
    ranked = sorted(
        sentences,
        key=lambda item: score_sentence(item, keywords),
        reverse=True,
    )
    selected = []
    seen = set()
    for sentence in ranked:
        normalized = re.sub(r"\s+", "", sentence)
        if normalized in seen:
            continue
        if score_sentence(sentence, keywords) <= 0:
            continue
        selected.append(sentence)
        seen.add(normalized)
        if len(selected) >= limit:
            break
    return selected


def contains_any(text, keywords):
    return any(keyword.lower() in text.lower() for keyword in keywords)


def trim_sentence(sentence, max_length=120):
    sentence = clean_text(sentence)
    if len(sentence) <= max_length:
        return sentence.rstrip("。；;，, .")
    return sentence[: max_length - 1] + "..."


def ensure_items(items, fallback_items, limit):
    result = list(items[:limit])
    for item in fallback_items:
        if len(result) >= limit:
            break
        if item not in result:
            result.append(item)
    return result[:limit]


def infer_view(sentences, asset_name):
    joined = "\n".join(sentences)
    positive_score = sum(joined.count(word) for word in POSITIVE_WORDS)
    negative_score = sum(joined.count(word) for word in NEGATIVE_WORDS)

    if asset_name == "债券":
        positive_score += joined.count("收益率下行") + joined.count("降息") + joined.count("久期机会")
        negative_score += joined.count("收益率上行") + joined.count("加息")
    if asset_name == "现金/货币基金":
        if contains_any(joined, ["短期", "防守", "流动性", "等待"]):
            positive_score += 2

    if positive_score >= negative_score + 2:
        return "看多"
    if negative_score >= positive_score + 2:
        return "谨慎"
    if negative_score > positive_score and negative_score >= 2:
        return "看空"
    return "中性"


def view_explanation(asset_name, view):
    explanations = {
        "看多": f"{asset_name}的正面线索更多，报告更强调机会和支撑因素。",
        "中性": f"{asset_name}机会和风险比较接近，更适合作为均衡配置而不是单边押注。",
        "谨慎": f"{asset_name}仍有配置价值，但报告提示的压力和波动不能忽视。",
        "看空": f"{asset_name}负面线索更集中，报告更强调回避或降低暴露。",
    }
    return explanations.get(view, explanations["中性"])


def find_focus(sentences, sub_assets, default_focus):
    found = []
    joined = "\n".join(sentences)
    for asset in sub_assets:
        if asset.lower() in joined.lower():
            found.append(asset)
    if found:
        return "、".join(found)
    return default_focus


def build_asset_view(asset_name, sentences, all_sentences, core_sentences, risks):
    config = ASSET_CONFIG[asset_name]
    asset_sentences = pick_sentences(all_sentences, config["keywords"], limit=8)
    context = asset_sentences or core_sentences or all_sentences[:5]

    view = infer_view(context, asset_name)
    reasons = ensure_items(
        [f"{trim_sentence(item)}。为什么：这是报告中直接提到的{asset_name}相关线索。"
         for item in context[:5]],
        [
            f"{view_explanation(asset_name, view)}为什么：报告中围绕宏观、政策、利率或盈利的表述会影响这类资产表现。",
            f"需要结合组合目标使用。为什么：同一类资产既可能承担收益来源，也可能承担分散风险的角色。",
            f"不宜只看短期涨跌。为什么：市场展望更关注未来一段时间的趋势和风险补偿。",
        ],
        3,
    )
    asset_risks = pick_sentences(asset_sentences + risks + all_sentences, RISK_KEYWORDS + config["keywords"], limit=4)
    risk_items = ensure_items(
        [f"{trim_sentence(item)}。为什么：这会影响客户持有体验或资产价格波动。"
         for item in asset_risks[:3]],
        [
            f"政策和数据变化可能改变{asset_name}的判断。为什么：市场会根据新的利率、通胀、增长数据重新定价。",
            f"短期波动可能放大客户情绪。为什么：即使中期逻辑没有改变，价格也可能先出现反复。",
        ],
        2,
    )

    return {
        "view": view,
        "focus": find_focus(context, config["sub_assets"], config["default_focus"]),
        "reasons": reasons[:3],
        "risks": risk_items[:2],
        "explain": build_client_explanation(asset_name, view, context),
    }


def build_client_explanation(asset_name, view, sentences):
    reason = trim_sentence(sentences[0], 100) if sentences else "报告没有给出足够细的单项资产描述"
    if asset_name == "股票":
        return f"可以把股票理解成看企业赚钱能力和市场情绪。当前判断为{view}，因为报告线索显示：{reason}。"
    if asset_name == "债券":
        return f"债券主要看利率方向。当前判断为{view}，因为报告线索显示：{reason}。利率往下时，存量债券通常更有支撑。"
    if asset_name == "外汇":
        return f"汇率本质上是在比较不同经济体的利率、增长和资金流向。当前判断为{view}，因为报告线索显示：{reason}。"
    if asset_name == "商品":
        return f"商品价格通常受供需、通胀和避险情绪影响。当前判断为{view}，因为报告线索显示：{reason}。"
    return f"现金和货币基金主要解决流动性问题。当前判断为{view}，因为报告线索显示：{reason}。"


def get_file_bytes(file):
    if hasattr(file, "getvalue"):
        return file.getvalue()
    return file.read()


def read_pdf(file):
    reader = PdfReader(BytesIO(get_file_bytes(file)))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def read_docx(file):
    document = Document(BytesIO(get_file_bytes(file)))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return "\n".join(paragraphs)


def read_txt(file):
    raw = get_file_bytes(file)
    for encoding in ("utf-8", "gb18030", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def image_metadata(image_bytes):
    try:
        image = Image.open(BytesIO(image_bytes))
        width, height = image.size
        fmt = (image.format or "PNG").lower()
    except UnidentifiedImageError:
        return None

    if width < MIN_CHART_WIDTH or height < MIN_CHART_HEIGHT:
        return None
    if width / max(height, 1) > 5 or height / max(width, 1) > 4:
        return None

    mime_type = "image/png" if fmt == "png" else "image/jpeg"
    return {
        "width": width,
        "height": height,
        "mime_type": mime_type,
    }


def chart_relevance_score(text):
    text = text or ""
    score = 0
    for keyword in CHART_RELEVANCE_KEYWORDS:
        if keyword.lower() in text.lower():
            score += 1
    if re.search(r"\d+(\.\d+)?%|\d{4}|[+-]\d", text):
        score += 2
    return score


def image_chart_score(metadata, context_text):
    score = chart_relevance_score(context_text)
    width = metadata["width"]
    height = metadata["height"]
    if width >= 600 and height >= 320:
        score += 1
    if width >= 900 and height >= 450:
        score += 1
    return score


def add_unique_image(images, seen_hashes, image_bytes, source, caption, context_text=""):
    digest = hashlib.sha256(image_bytes).hexdigest()
    if digest in seen_hashes:
        return

    metadata = image_metadata(image_bytes)
    if metadata is None:
        return
    score = image_chart_score(metadata, context_text)
    if score < 3:
        return

    images.append(
        {
            "bytes": image_bytes,
            "source": source,
            "caption": caption,
            "width": metadata["width"],
            "height": metadata["height"],
            "mime_type": metadata["mime_type"],
            "score": score,
            "context": trim_sentence(context_text, 160),
        }
    )
    seen_hashes.add(digest)


def extract_pdf_images(file):
    images = []
    seen_hashes = set()
    pdf_bytes = get_file_bytes(file)
    document = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page_index in range(len(document)):
        page = document[page_index]
        page_text = page.get_text("text") or ""
        previous_text = document[page_index - 1].get_text("text") if page_index > 0 else ""
        next_text = document[page_index + 1].get_text("text") if page_index + 1 < len(document) else ""
        context_text = clean_text("\n".join([previous_text[-600:], page_text, next_text[:600]]))
        for image_index, image_info in enumerate(page.get_images(full=True), start=1):
            if len(images) >= MAX_EXTRACTED_IMAGES:
                return sorted(images, key=lambda item: item["score"], reverse=True)
            xref = image_info[0]
            try:
                extracted = document.extract_image(xref)
            except Exception:
                continue
            image_bytes = extracted.get("image")
            if not image_bytes:
                continue
            caption = f"{file.name} - 第 {page_index + 1} 页图表/图片 {image_index}"
            add_unique_image(images, seen_hashes, image_bytes, file.name, caption, context_text)

    return sorted(images, key=lambda item: item["score"], reverse=True)


def extract_docx_images(file):
    images = []
    seen_hashes = set()
    docx_bytes = get_file_bytes(file)
    document = Document(BytesIO(docx_bytes))
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    relevant_document = chart_relevance_score(document_text) >= 4
    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        media_files = [
            name for name in archive.namelist()
            if name.startswith("word/media/")
        ]
        for image_index, name in enumerate(media_files, start=1):
            if len(images) >= MAX_EXTRACTED_IMAGES:
                return sorted(images, key=lambda item: item["score"], reverse=True)
            image_bytes = archive.read(name)
            caption = f"{file.name} - 图表/图片 {image_index}"
            context_text = document_text if relevant_document else ""
            add_unique_image(images, seen_hashes, image_bytes, file.name, caption, context_text)
    return sorted(images, key=lambda item: item["score"], reverse=True)


def extract_uploaded_images(file):
    name = file.name.lower()
    try:
        if name.endswith(".pdf"):
            return extract_pdf_images(file)
        if name.endswith(".docx"):
            return extract_docx_images(file)
    except Exception as exc:
        st.warning(f"{file.name} 图表提取失败：{exc}")
    return []


def read_uploaded_file(file):
    name = file.name.lower()
    if name.endswith(".pdf"):
        return read_pdf(file)
    if name.endswith(".docx"):
        return read_docx(file)
    if name.endswith(".txt"):
        return read_txt(file)
    return ""


def collect_input_documents(uploaded_files, url_text, pasted_text):
    documents = []

    if uploaded_files:
        if len(uploaded_files) > MAX_COMPARE_DOCUMENTS:
            st.error(f"最多支持上传 {MAX_COMPARE_DOCUMENTS} 份报告做对比，请减少文件数量后再生成。")
            st.stop()

        for uploaded_file in uploaded_files:
            try:
                text = read_uploaded_file(uploaded_file)
                if text:
                    documents.append({"name": uploaded_file.name, "text": text})
            except Exception as exc:
                st.warning(f"{uploaded_file.name} 读取失败：{exc}")

    urls = [line.strip() for line in url_text.splitlines() if line.strip()]
    for url in urls:
        with st.spinner(f"正在读取网页：{url}"):
            text, error = fetch_url_text(url)
        if text:
            documents.append({"name": f"网页链接：{url}", "text": text})
            st.success(f"已读取网页：{url}")
        else:
            st.warning(f"{url} 读取失败：{error}。请手动复制网页正文，粘贴到右侧输入框。")

    if pasted_text.strip():
        documents.append({"name": "手动粘贴内容", "text": pasted_text})

    all_text = clean_text("\n\n".join(document["text"] for document in documents))
    return documents, all_text


def normalize_url(url):
    url = url.strip()
    if not url:
        return ""
    parsed = urlparse(url)
    if not parsed.scheme:
        url = "https://" + url
        parsed = urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        return ""
    return url


def extract_text_from_html(html):
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
        tag.decompose()

    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    description_tag = soup.find("meta", attrs={"name": "description"})
    description = ""
    if description_tag:
        description = description_tag.get("content", "").strip()

    containers = soup.find_all(["article", "main"])
    if not containers:
        containers = [soup]

    paragraphs = []
    for container in containers:
        for tag in container.find_all(["h1", "h2", "h3", "p", "li"]):
            text = clean_text(tag.get_text(" ", strip=True))
            if len(text) >= 20:
                paragraphs.append(text)

    unique_paragraphs = []
    seen = set()
    for paragraph in paragraphs:
        normalized = re.sub(r"\s+", "", paragraph)
        if normalized in seen:
            continue
        unique_paragraphs.append(paragraph)
        seen.add(normalized)

    parts = [title, description] + unique_paragraphs
    return clean_text("\n".join(part for part in parts if part))


def fetch_url_text(url):
    normalized_url = normalize_url(url)
    if not normalized_url:
        return "", "链接格式不正确"

    try:
        response = requests.get(
            normalized_url,
            headers=REQUEST_HEADERS,
            timeout=12,
            allow_redirects=True,
        )
        response.raise_for_status()
    except requests.RequestException as exc:
        return "", f"网页访问失败：{exc}"

    content_type = response.headers.get("content-type", "").lower()

    if "pdf" in content_type or normalized_url.lower().endswith(".pdf"):
        try:
            reader = PdfReader(BytesIO(response.content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            text = clean_text(text)
        except Exception as exc:
            return "", f"PDF 链接读取失败：{exc}"
    else:
        response.encoding = response.apparent_encoding or response.encoding
        text = extract_text_from_html(response.text)

    if len(text) < 80:
        return "", "读取到的正文太短，可能被网站限制、需要登录，或页面主要由脚本加载"

    return text, ""


def bullet_list(items, empty_text):
    if not items:
        return f"- {empty_text}"
    return "\n".join(f"- {item}" for item in items)


def first_item(items, fallback):
    if items:
        return items[0]
    return fallback


def build_environment(sentences):
    environment = []
    for label, keywords in ENVIRONMENT_LABELS.items():
        picked = pick_sentences(sentences, keywords, limit=1)
        if picked:
            environment.append(f"{label}：{trim_sentence(picked[0])}。为什么：这是影响各类资产定价的关键变量。")
    return ensure_items(
        environment,
        [
            "风险偏好：需要观察市场是否愿意承担波动。为什么：风险偏好回升通常利好权益，下降则更利好防守资产。",
            "利率周期：需要关注央行政策和收益率变化。为什么：利率决定债券价格，也会影响股票估值。",
            "美元趋势：需要关注美元强弱。为什么：美元会影响人民币、黄金和新兴市场资产表现。",
        ],
        3,
    )


def build_strategy(asset_views):
    overweights = [name for name, item in asset_views.items() if item["view"] == "看多"]
    underweights = [name for name, item in asset_views.items() if item["view"] in ("谨慎", "看空")]
    neutral_assets = [name for name, item in asset_views.items() if item["view"] == "中性"]

    if not overweights:
        overweights = neutral_assets[:2] or ["债券", "现金/货币基金"]
    if not underweights:
        underweights = ["波动较大、报告证据不足的资产"]

    core = []
    if asset_views["债券"]["view"] in ("看多", "中性"):
        core.append("债券")
    if asset_views["现金/货币基金"]["view"] in ("看多", "中性"):
        core.append("现金/货币基金")
    if asset_views["股票"]["view"] in ("看多", "中性"):
        core.append("优质股票资产")
    core = core or ["现金/货币基金", "中低波动资产"]

    opportunities = []
    if asset_views["股票"]["view"] in ("看多", "中性"):
        opportunities.append("股票")
    if asset_views["商品"]["view"] in ("看多", "中性"):
        opportunities.append("黄金/商品")
    if asset_views["外汇"]["view"] in ("看多", "中性"):
        opportunities.append("美元或其他外币资产")
    opportunities = opportunities or ["报告中有明确催化剂的单项资产"]

    controlled = [name for name, item in asset_views.items() if item["view"] in ("谨慎", "看空")]
    controlled = controlled or ["短期涨幅过大、客户承受不了波动的资产"]

    return {
        "overweights": overweights,
        "underweights": underweights,
        "core": core,
        "opportunities": opportunities,
        "controlled": controlled,
    }


def build_wechat_outlook(core, environment, asset_views):
    stock_view = asset_views["股票"]["view"]
    bond_view = asset_views["债券"]["view"]
    commodity_view = asset_views["商品"]["view"]
    main = trim_sentence(core[0], 70) if core else "近期市场围绕利率、政策和增长预期反复定价"
    env = environment[0].split("。为什么：")[0] if environment else "市场仍在观察经济数据和政策节奏"
    return (
        f"近期市场主线是{main}。背后原因主要是{env}，资金在股票、债券和商品之间重新权衡。"
        f"五月看，股票整体偏{stock_view}，债券偏{bond_view}，商品偏{commodity_view}。"
        f"配置上建议更重视资产分散，把稳健底仓和阶段性机会分开看，不因单日波动做大幅调整。"
    )


def build_phone_script(core, environment, asset_views, strategy):
    main = trim_sentence(core[0], 100) if core else "这份展望的主线是市场正在重新评估增长、通胀和政策节奏"
    env = environment[0].split("。为什么：")[0] if environment else "当前市场还在看政策和经济数据"
    return (
        f"您好，我简单跟您同步一下这份五月市场展望。核心不是说某个资产马上会涨，"
        f"而是市场现在主要围绕一条线在交易：{main}。{env}，所以不同资产的表现会分化。"
        f"股票这边报告给出的判断偏{asset_views['股票']['view']}，关键还是看盈利和估值能不能继续配合；"
        f"债券偏{asset_views['债券']['view']}，主要看后面利率和降息预期怎么走；"
        f"外汇要关注美元和人民币的相对变化，商品里黄金、原油、铜也会受到避险、供需和通胀影响。"
        f"对配置的启发是，不建议只押单一方向，可以把{ '、'.join(strategy['core'][:2]) }作为底仓，"
        f"再把{ '、'.join(strategy['opportunities'][:2]) }作为机会型观察。这样组合既能参与机会，也能控制波动。"
    )


def build_client_questions(asset_views):
    return [
        (
            "现在还能买美股吗？美股都涨了这么高。",
            f"可以关注，但不适合只因为上涨就追高。为什么：股票判断偏{asset_views['股票']['view']}，后续更取决于盈利能否兑现、估值是否合理，以及利率是否继续压制估值。客户如果已有较高仓位，应更重视分批和控制比例。",
        ),
        (
            "港股跌了这么多，还能不能加仓？",
            "港股可以作为机会型资产观察，但要看政策、盈利和资金流是否改善。为什么：便宜不等于马上上涨，只有当基本面或资金面出现改善，低估值才更容易转化为回报。",
        ),
        (
            "黄金还适合配置吗？",
            f"黄金更适合作为组合里的分散和防守资产，不宜当成短线投机。为什么：商品判断偏{asset_views['商品']['view']}，黄金通常受避险需求、实际利率和美元强弱影响，能对冲一部分市场不确定性。",
        ),
        (
            "美元资产还要不要留？",
            f"美元资产可以保留一定比例，但要看客户本身是否有外币需求。为什么：外汇判断偏{asset_views['外汇']['view']}，美元强弱会受美国利率、降息节奏和全球风险偏好影响，单纯赌汇率并不适合多数客户。",
        ),
        (
            "债券现在有没有机会？",
            f"债券值得关注，尤其适合做组合稳定器。为什么：债券判断偏{asset_views['债券']['view']}，如果后续利率下行，债券价格通常更有支撑；但如果通胀反复或降息推迟，债券也会有波动。",
        ),
    ]


def summarize_market(text):
    sentences = split_sentences(text)

    core = pick_sentences(sentences, CORE_KEYWORDS, limit=4)
    risks = pick_sentences(sentences, RISK_KEYWORDS, limit=4)

    if not core:
        core = sentences[:3]

    environment = build_environment(sentences)
    asset_views = {}
    for asset_name in ASSET_CONFIG:
        asset_views[asset_name] = build_asset_view(asset_name, sentences, sentences, core, risks)

    strategy = build_strategy(asset_views)
    wechat = build_wechat_outlook(core, environment, asset_views)
    phone = build_phone_script(core, environment, asset_views, strategy)
    questions = build_client_questions(asset_views)
    report = build_full_report(core, environment, asset_views, strategy, wechat, phone, questions)

    return {
        "core": core,
        "environment": environment,
        "asset_views": asset_views,
        "risks": risks,
        "strategy": strategy,
        "wechat": wechat,
        "phone": phone,
        "questions": questions,
        "report": report,
    }


def markdown_list(items):
    return "\n".join(f"- {item}" for item in items)


def build_full_report(core, environment, asset_views, strategy, wechat, phone, questions):
    main_line = trim_sentence(core[0], 120) if core else "本文主线是市场围绕增长、利率、美元和政策预期重新定价。"
    important = ensure_items(
        [f"{trim_sentence(item)}。为什么：这是报告中反复影响资产判断的关键线索。" for item in core[:5]],
        [
            "资产表现会出现分化。为什么：股票、债券、外汇和商品分别受盈利、利率、美元和供需影响。",
            "组合需要兼顾进攻和防守。为什么：单一资产很难同时处理收益机会和短期波动。",
            "客户沟通要把市场主线讲清楚。为什么：客户真正关心的是为什么涨跌、接下来如何配置。",
        ],
        3,
    )

    lines = [
        "## 一、这份市场展望的核心结论",
        "",
        f"1. 一句话主线：{main_line}",
        "",
        "2. 重要观点：",
        markdown_list(important),
        "",
        "3. 当前市场环境：",
        markdown_list(environment),
        "",
        "## 二、各资产大类未来观点与依据",
        "",
    ]

    lines.extend(build_stock_section(asset_views["股票"]))
    lines.extend(build_bond_section(asset_views["债券"]))
    lines.extend(build_fx_section(asset_views["外汇"]))
    lines.extend(build_commodity_section(asset_views["商品"]))
    lines.extend(build_cash_section(asset_views["现金/货币基金"]))

    lines.extend(
        [
            "## 三、投资策略总结",
            "",
            f"1. 应该超配什么：{'、'.join(strategy['overweights'])}。为什么：这些资产在报告里体现出的机会线索相对更多，适合作为阶段性重点。",
            f"2. 应该低配什么：{'、'.join(strategy['underweights'])}。为什么：这些资产的风险、波动或证据不足更突出，不适合作为组合重心。",
            f"3. 适合做核心持仓：{'、'.join(strategy['core'])}。为什么：核心持仓要承担长期配置和控制波动的作用，不能只追短期弹性。",
            f"4. 适合做机会型配置：{'、'.join(strategy['opportunities'])}。为什么：这些资产可能受政策、利率、资金或供需变化带动，但更需要控制节奏。",
            f"5. 需要控制仓位：{'、'.join(strategy['controlled'])}。为什么：如果客户承受不了波动，仓位过高会影响持有体验和决策稳定性。",
            "",
            "## 四、给客户的五月市场展望微信版",
            "",
            wechat,
            "",
            "## 五、电话汇报版",
            "",
            phone,
            "",
            "## 六、客户可能会问的问题",
            "",
        ]
    )

    for index, (question, answer) in enumerate(questions, start=1):
        lines.append(f"{index}. {question}")
        lines.append(f"答：{answer}")
        lines.append("")

    return "\n".join(lines)


def build_stock_section(item):
    return [
        "### 1. 股票",
        f"- 未来观点：{item['view']}。为什么：{view_explanation('股票', item['view'])}",
        f"- 重点区域：{item['focus']}。",
        "- 支撑理由：",
        markdown_list(item["reasons"]),
        "- 主要风险：",
        markdown_list(item["risks"]),
        f"- 客户可理解的解释：{item['explain']}",
        "",
    ]


def build_bond_section(item):
    return [
        "### 2. 债券",
        f"- 未来观点：{item['view']}。为什么：{view_explanation('债券', item['view'])}",
        f"- 利率/降息/收益率曲线判断：{item['focus']}。为什么：债券价格主要受利率和收益率变化影响，降息预期越明确，债券越容易获得支撑。",
        "- 支撑理由：",
        markdown_list(item["reasons"]),
        "- 主要风险：",
        markdown_list(item["risks"]),
        f"- 客户可理解的解释：{item['explain']}",
        "",
    ]


def build_fx_section(item):
    return [
        "### 3. 外汇",
        f"- 美元、人民币、日元、欧元等主要判断：{item['view']}，重点观察{item['focus']}。为什么：汇率由利率差、经济强弱和资金流向共同决定。",
        "- 支撑理由：",
        markdown_list(item["reasons"]),
        "- 主要风险：",
        markdown_list(item["risks"]),
        f"- 客户可理解的解释：{item['explain']}",
        "",
    ]


def build_commodity_section(item):
    return [
        "### 4. 商品",
        f"- 黄金、原油、铜等主要判断：{item['view']}，重点观察{item['focus']}。为什么：黄金偏避险，原油和铜更受供需与增长预期影响。",
        "- 支撑理由：",
        markdown_list(item["reasons"]),
        "- 主要风险：",
        markdown_list(item["risks"]),
        f"- 客户可理解的解释：{item['explain']}",
        "",
    ]


def build_cash_section(item):
    return [
        "### 5. 现金/货币基金",
        f"- 是否适合作为短期配置：{item['view']}。为什么：现金和货币基金主要提供流动性，适合承接短期不用但又不想承受太大波动的资金。",
        f"- 在当前市场下的角色：{item['focus']}。为什么：当市场方向不够清晰时，现金类资产可以给客户保留后续加仓或再平衡空间。",
        "- 支撑理由：",
        markdown_list(item["reasons"]),
        "- 主要风险：",
        markdown_list(item["risks"]),
        f"- 客户可理解的解释：{item['explain']}",
        "",
    ]


def build_wechat_version(core, sections, risks):
    stock = sections.get("股票", [])
    bond = sections.get("债券", [])
    fx = sections.get("外汇", [])
    commodity = sections.get("商品", [])

    lines = [
        "【今日市场简报】",
        "",
        "核心观点：",
        bullet_list(core[:3], "今日市场信息有限，建议补充更多原文后再生成。"),
        "",
        "分板块看：",
        f"- 股票：{stock[0] if stock else '暂无明确股票市场信息。'}",
        f"- 债券：{bond[0] if bond else '暂无明确债券市场信息。'}",
        f"- 外汇：{fx[0] if fx else '暂无明确外汇市场信息。'}",
        f"- 商品：{commodity[0] if commodity else '暂无明确商品市场信息。'}",
        "",
        "风险提示：",
        bullet_list(risks[:3], "需关注宏观数据、政策变化和市场波动。"),
    ]
    return "\n".join(lines)


def build_phone_version(core, sections, risks):
    core_text = "；".join(core[:2]) if core else "今天的核心信息还不够充分"
    stock = first_item(sections.get("股票", []), "股票方面暂无特别明确的信息")
    bond = first_item(sections.get("债券", []), "债券方面暂无特别明确的信息")
    fx = first_item(sections.get("外汇", []), "外汇方面暂无特别明确的信息")
    commodity = first_item(sections.get("商品", []), "商品方面暂无特别明确的信息")
    risk_text = "；".join(risks[:2]) if risks else "主要还是提醒客户注意数据、政策和短期波动"

    return (
        f"今天可以这样和客户口头汇报：整体来看，{core_text}。"
        f"分板块看，股票方面，{stock}；债券方面，{bond}；"
        f"外汇方面，{fx}；商品方面，{commodity}。"
        f"风险上，{risk_text}。整体建议保持跟踪，不要只看单日波动。"
    )


def render_result(result):
    read_tab, copy_tab = st.tabs(["阅读版", "复制版"])

    with read_tab:
        st.subheader("市场观点分析")
        st.markdown(result)

    with copy_tab:
        st.subheader("可复制完整版本")
        st.text_area("完整市场观点分析", result, height=620)


def render_copy_panel(title, text, key):
    text = text or ""
    safe_title = escape(title)
    safe_text = escape(text)
    text_json = json.dumps(text, ensure_ascii=False).replace("</", "<\\/")
    safe_key = re.sub(r"[^a-zA-Z0-9_]", "_", key)
    height = max(220, min(430, 150 + len(text) // 2))

    components.html(
        f"""
        <div style="font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;border:1px solid #ead4d4;border-radius:14px;padding:18px 18px 16px;background:#fffafa;">
            <div style="display:flex;align-items:center;justify-content:space-between;gap:12px;margin-bottom:12px;">
                <div style="font-weight:700;color:#1f2937;font-size:16px;">{safe_title}</div>
                <button id="copy-{safe_key}" style="border:1px solid #c8102e;background:#c8102e;color:white;border-radius:999px;padding:8px 14px;font-weight:700;cursor:pointer;">📋 复制</button>
            </div>
            <div style="white-space:pre-wrap;line-height:1.75;color:#20252d;font-size:15px;background:white;border:1px solid #f2dddd;border-radius:12px;padding:14px;">{safe_text}</div>
            <div id="msg-{safe_key}" style="display:none;margin-top:10px;color:#c8102e;font-size:13px;font-weight:700;">已复制</div>
        </div>
        <script>
        const copyText_{safe_key} = {text_json};
        const button_{safe_key} = document.getElementById("copy-{safe_key}");
        const message_{safe_key} = document.getElementById("msg-{safe_key}");
        button_{safe_key}.onclick = async () => {{
            try {{
                await navigator.clipboard.writeText(copyText_{safe_key});
            }} catch (error) {{
                const area = document.createElement("textarea");
                area.value = copyText_{safe_key};
                document.body.appendChild(area);
                area.select();
                document.execCommand("copy");
                document.body.removeChild(area);
            }}
            message_{safe_key}.style.display = "block";
            setTimeout(() => message_{safe_key}.style.display = "none", 1600);
        }};
        </script>
        """,
        height=height,
        scrolling=True,
    )


def save_touch_history(result):
    entry_id = hashlib.sha256(
        json.dumps(result, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()[:12]
    entry = {
        "id": entry_id,
        "one_liner": result.get("one_liner", ""),
        "deep_push": result.get("deep_push", ""),
        "care_touch": result.get("care_touch", ""),
    }
    entry_json = json.dumps(entry, ensure_ascii=False).replace("</", "<\\/")

    components.html(
        f"""
        <script>
        const storageKey = "marketBriefingTouchHistory";
        const entry = {entry_json};
        const now = new Date();
        entry.time = now.toLocaleString("zh-CN", {{ hour12: false }});
        const existing = JSON.parse(localStorage.getItem(storageKey) || "[]");
        const updated = [entry, ...existing.filter(item => item.id !== entry.id)].slice(0, 10);
        localStorage.setItem(storageKey, JSON.stringify(updated));
        </script>
        """,
        height=0,
    )


def render_touch_history():
    components.html(
        """
        <div style="font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;border-top:1px solid #ead4d4;padding-top:14px;">
            <div style="font-weight:800;color:#1f2937;margin-bottom:10px;">最近 10 次生成历史</div>
            <div id="touch-history"></div>
        </div>
        <script>
        const storageKey = "marketBriefingTouchHistory";
        const box = document.getElementById("touch-history");
        const items = JSON.parse(localStorage.getItem(storageKey) || "[]");
        function escapeHtml(value) {
            return String(value || "").replace(/[&<>"']/g, s => ({
                "&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;", "'": "&#39;"
            }[s]));
        }
        async function copyText(text, id) {
            try {
                await navigator.clipboard.writeText(text);
            } catch (error) {
                const area = document.createElement("textarea");
                area.value = text;
                document.body.appendChild(area);
                area.select();
                document.execCommand("copy");
                document.body.removeChild(area);
            }
            const tip = document.getElementById(`history-tip-${id}`);
            if (tip) {
                tip.style.display = "inline";
                setTimeout(() => tip.style.display = "none", 1400);
            }
        }
        if (!items.length) {
            box.innerHTML = '<div style="color:#6b7280;font-size:14px;">生成后会自动保存在本机浏览器里。</div>';
        } else {
            box.innerHTML = items.map((item, index) => {
                const allText = `🎯 一句话市场观点\\n${item.one_liner || ""}\\n\\n📊 针对性深度推送\\n${item.deep_push || ""}\\n\\n💝 钩子型触达(撩客户)\\n${item.care_touch || ""}`;
                window[`historyCopyText_${index}`] = allText;
                return `
                    <details style="border:1px solid #f0dddd;border-radius:12px;background:#fffafa;margin-bottom:10px;padding:10px;">
                        <summary style="cursor:pointer;color:#20252d;font-weight:700;">${escapeHtml(item.time)}｜${escapeHtml(item.one_liner).slice(0, 42)}</summary>
                        <div style="white-space:pre-wrap;line-height:1.65;color:#374151;font-size:13px;margin:10px 0;">${escapeHtml(allText)}</div>
                        <button onclick="copyText(window.historyCopyText_${index}, ${index})" style="border:1px solid #c8102e;background:white;color:#c8102e;border-radius:999px;padding:6px 12px;font-weight:700;cursor:pointer;">📋 复制这次</button>
                        <span id="history-tip-${index}" style="display:none;color:#c8102e;font-size:12px;margin-left:8px;font-weight:700;">已复制</span>
                    </details>
                `;
            }).join("");
        }
        </script>
        """,
        height=380,
        scrolling=True,
    )


def render_touch_copy_result(result):
    st.subheader("客户触达文案")
    touch_tabs = st.tabs([
        "🎯 一句话市场观点",
        "📊 针对性深度推送",
        "💝 钩子型触达(撩客户)",
    ])

    with touch_tabs[0]:
        render_copy_panel("一句话市场观点（用于周一群发）", result["one_liner"], "one-liner")
    with touch_tabs[1]:
        render_copy_panel("针对性深度推送", result["deep_push"], "deep-push")
    with touch_tabs[2]:
        render_copy_panel("钩子型触达(撩客户)", result["care_touch"], "care-touch")
        st.caption("目标:让不关心市场的客户也想回复你")

    save_touch_history(result)
    render_touch_history()


def render_followup_chat(model_name):
    if not st.session_state.get("analysis_result"):
        return

    st.divider()
    st.subheader("继续追问")
    st.caption("可以基于本次材料继续追问，也可以让 DeepSeek 结合通用市场知识补充判断；回答会区分原文依据和模型补充。")

    col_a, col_b = st.columns([1, 4])
    with col_a:
        if st.button("清空追问记录"):
            st.session_state["followup_messages"] = []
            st.rerun()

    for message in st.session_state.get("followup_messages", []):
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    question = st.chat_input("输入你的追问")
    if question:
        st.session_state.setdefault("followup_messages", [])
        st.session_state["followup_messages"].append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)

        try:
            with st.chat_message("assistant"):
                with st.spinner("正在结合材料和模型知识继续分析..."):
                    answer = call_deepseek_followup(question, model_name.strip() or DEFAULT_DEEPSEEK_MODEL)
                st.markdown(answer)
            st.session_state["followup_messages"].append({"role": "assistant", "content": answer})
        except Exception as exc:
            st.error(f"追问失败：{exc}")


api_ready = bool(os.getenv("DEEPSEEK_API_KEY"))
api_status_class = "ready" if api_ready else "warn"
api_status_text = "DeepSeek 已配置" if api_ready else "等待配置 API Key"

st.markdown(
    f"""
    <div class="app-hero">
        <div class="app-kicker">财富管理客户沟通工作台</div>
        <h1 class="app-title">Jack 市场沟通助手</h1>
        <p class="app-subtitle">
            上传市场展望、研报或财经网页内容，生成客户经理可直接使用的市场观点分析、短期机会提示、图表线索和客户沟通话术。
        </p>
        <div class="status-row">
            <span class="status-pill {api_status_class}">{api_status_text}</span>
            <span class="status-pill">PDF / Word / TXT</span>
            <span class="status-pill">最多 5 份报告对比</span>
            <span class="status-pill">网页链接读取</span>
            <span class="status-pill">图表线索提示</span>
            <span class="status-pill">三版客户触达文案</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

st.subheader("输入材料")

left, right = st.columns([1, 1])

with left:
    st.markdown("#### 文件和链接")
    uploaded_files = st.file_uploader(
        "上传 PDF、Word、TXT 文件（最多 5 份报告）",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
    )
    url_text = st.text_area(
        "粘贴财经新闻/报告链接",
        placeholder="一行一个链接，例如：https://example.com/market-news",
        height=140,
    )

with right:
    st.markdown("#### 模型和正文")
    model_name = st.text_input("DeepSeek 模型名称", value=DEFAULT_DEEPSEEK_MODEL)
    output_style = st.selectbox(
        "输出风格",
        options=list(OUTPUT_STYLES.keys()),
        index=0,
        help="选择材料最终给谁看，DeepSeek 会按对应受众调整解释深度、语言和重点。",
    )
    pasted_text = st.text_area(
        "粘贴财经网站内容",
        placeholder="把你从财经网站复制来的市场评论、新闻或研报摘要粘贴到这里。",
        height=220,
    )

st.divider()

button_left, button_right = st.columns([1, 1])
with button_left:
    generate_analysis = st.button("生成市场观点分析", type="primary", use_container_width=True)
with button_right:
    generate_touch_copy = st.button("生成客户触达文案", use_container_width=True)

if generate_analysis or generate_touch_copy:
    documents, all_text = collect_input_documents(uploaded_files, url_text, pasted_text)

    if not all_text:
        st.error("请先上传文件，或粘贴一段市场内容。")
    else:
        if generate_analysis:
            try:
                if len(documents) >= 2:
                    st.info(f"已识别 {len(documents)} 份材料，将生成多文档对比分析，包括共识观点和分歧点。")
                with st.spinner("正在调用 DeepSeek 生成市场观点分析..."):
                    result = call_deepseek_market_analysis(
                        all_text,
                        model_name.strip() or DEFAULT_DEEPSEEK_MODEL,
                        output_style,
                        documents,
                    )
                st.session_state["analysis_source_text"] = all_text
                st.session_state["analysis_result"] = result
                st.session_state["followup_messages"] = []
                st.divider()
                render_result(result)
            except Exception as exc:
                st.error(f"DeepSeek 调用失败：{exc}")
                st.info("请检查 .env 里的 DEEPSEEK_API_KEY 是否正确、网络是否可用、模型名称是否有效。")

        if generate_touch_copy:
            try:
                if len(documents) >= 2:
                    st.info(f"已识别 {len(documents)} 份材料，将综合提炼成客户触达文案。")
                with st.spinner("正在调用 DeepSeek 生成 3 个客户触达版本..."):
                    touch_result = call_deepseek_touch_copy(
                        all_text,
                        model_name.strip() or DEFAULT_DEEPSEEK_MODEL,
                    )
                st.session_state["touch_copy_result"] = touch_result
                st.divider()
                render_touch_copy_result(touch_result)
            except Exception as exc:
                st.error(f"客户触达文案生成失败：{exc}")
                st.info("请检查 .env 里的 DEEPSEEK_API_KEY 是否正确、网络是否可用、模型名称是否有效。")

if st.session_state.get("touch_copy_result") and not generate_touch_copy:
    st.divider()
    render_touch_copy_result(st.session_state["touch_copy_result"])

render_followup_chat(model_name)

with st.sidebar:
    st.header("使用说明")
    st.write("1. 上传 PDF、Word 或 TXT 文件。")
    st.write("2. 可同时上传最多 5 份报告，DeepSeek 会输出共识观点和分歧点。")
    st.write("3. 也可以粘贴财经新闻/报告链接。")
    st.write("4. 工具不再自动提取图片；如果材料文字里提到关键图表，会提示你回原文查看。")
    st.write("5. 如果链接读取失败，就手动复制正文粘贴进输入框。")
    st.write("6. 可以选择略懂客户、小白客户、客户经理自用或内部群同事分享风格。")
    st.write("7. 在线部署时建议设置 APP_PASSWORD，避免他人随意消耗 API 额度。")
    st.write("8. 点击“生成市场观点分析”可以得到完整解读。")
    st.write("9. 点击“生成客户触达文案”可以得到 3 个可复制版本，并自动保留最近 10 次历史。")
    st.write("10. 生成分析后可以继续追问，DeepSeek 会结合本次材料和通用市场知识回答。")
    st.write("11. 复制完整报告，并按需回到原文核对关键图表或数据。")
    st.divider()
    st.write("说明：本版本通过 OpenAI Python SDK 的兼容方式调用 DeepSeek。")
